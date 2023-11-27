# import copy

import json

import azure.functions as func
import docx
from bs4 import BeautifulSoup
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# from docx.text.run import Run


def add_comment_to_table(table: Table, substring: str, comment: str):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if substring in run.text:
                        run.add_comment(comment, author="GPT")
                        # print(f"added comment '{comment_text}' to table cell")


def add_comment_to_paragraph(paragraph: Paragraph, substring: str, comment: str):
    # place = -1
    # new_runs: list[Run] = []
    for i, run in enumerate(paragraph.runs):
        if substring in run.text:
            # print(f"found run: {run.text}")
            # start = str(run.text).index(substring)
            # end = start + len(substring)
            # print("from", start)
            # print("to", end)
            # print(run.text[start:end])
            # print([run.text[:start], run.text[start:end], run.text[end:]])

            # run1 = copy.deepcopy(run)
            # run2 = copy.deepcopy(run)
            # run3 = copy.deepcopy(run)

            # run1.clear().add_text(run.text[:start])
            # run2.clear().add_text(run.text[start:end])
            # run3.clear().add_text(run.text[end:])

            run.add_comment(comment, "GPT")

            # new_runs = [run1, run2, run3]

            # print([r.text for r in [run1, run2, run3]])

            # paragraph.runs[i : i + 1] = [run2]
            # place = i
            # print("all runs:", [r.text for r in paragraph.runs])

            break

    # print(place)
    # print([r.text for r in new_runs])
    # print([r.text for r in paragraph.runs[place : place + 1]])
    # paragraph.runs[place : place + 1] = new_runs

    # backup = copy.deepcopy(paragraph.runs)
    # backup[place : place + 1] = new_runs

    # paragraph.runs.clear()
    # paragraph.append_runs(backup)

    # for run in backup:

    # paragraph.runs.insert(place, new_runs[1])
    # print([r.text for r in paragraph.runs])

    #     part1 = run.text[: run.text.index(substring)]
    #     part3 = run.text[len(part1) + len(substring) :]

    #     run1 = copy.deepcopy(run)
    #     run2 = copy.deepcopy(run)
    #     run3 = copy.deepcopy(run)

    #     run1.text = part1
    #     run2.text = substring
    #     run3.text = part3

    #     print("new run 1", run1.text)
    #     print("new run 2", run2.text)
    #     print("new run 3", run3.text)

    #     # run2.add_comment(comment, "GPT")

    #     print([r.text for r in paragraph.runs])

    #     paragraph.runs[i : i + 1] = []

    #     print([r.text for r in paragraph.runs])

    #     print(f"added comment '{comment_text}' to paragraph")


# def main(req: func.HttpRequest) -> func.HttpResponse:
def main():
    # payload = req.get_json()
    # html = payload["suggestions"]

    with open("suggestions.html", "r", encoding="utf-8") as html:
        try:
            doc: Document = docx.Document("no comments.docx")
            soup = BeautifulSoup(html, "html.parser")

            # print([f"{i}: {p.text.strip()}" for i, p in enumerate(doc.paragraphs)])

            suggestions = soup.select("div")

            for s in suggestions:
                index = int(str(s.get("data-id")))
                search_text = str(s.get("data-text")).strip()
                element_type = str(s.get("data-type")).strip()
                comment_text = s.get_text().strip()

                full_comment = f"{search_text}\n\n{comment_text}"

                if element_type == "paragraph":
                    add_comment_to_paragraph(
                        doc.paragraphs[index], search_text, full_comment
                    )

                elif element_type == "table_cell":
                    add_comment_to_table(doc.tables[index], search_text, full_comment)

            doc.save("yes comments.docx")

            response = {}

            return func.HttpResponse(json.dumps(response), status_code=200)
        except:
            return func.HttpResponse(
                json.dumps({"message": "Failed to process request"}), status_code=400
            )
